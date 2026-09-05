import streamlit as st
import requests
import datetime
import pandas as pd
import json
import re
from docx import Document
import io

def vygeneruj_formular_docx(nadpis, radky):
    doc = Document()
    doc.add_heading(nadpis, level=1)
    doc.add_paragraph("Metodika M-TECH CORE — oficiální formulář projektu")
    doc.add_paragraph("")
    for label, hodnota in radky:
        p = doc.add_paragraph()
        p.add_run(f"{label}: ").bold = True
        p.add_run(str(hodnota))
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

st.set_page_config(page_title="Startup Hub a Dashboard", layout="wide")

st.markdown("""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Montserrat:wght@400;600;800&display=swap');
    html, body, [class*="css"] { font-family: 'Montserrat', sans-serif !important; }
    h1, h2, h3, h4 { background: -webkit-linear-gradient(45deg, #00B4D8, #0077B6); -webkit-background-clip: text; -webkit-text-fill-color: transparent; font-weight: 800 !important; }
    .stButton>button { border-radius: 8px; transition: all 0.3s ease; border: 1px solid #00B4D8; width: 100%; font-weight: 600; }
    .stButton>button:hover { transform: translateY(-2px); box-shadow: 0 4px 15px rgba(0, 180, 216, 0.4); border-color: #00B4D8; background-color: #0f172a; color: white;}
    .card-box { background-color: #1e293b; padding: 20px; border-radius: 12px; border: 1px solid #334155; margin-bottom: 15px; }
    .status-badge-ok { background: rgba(16, 185, 129, 0.15); color: #34d399; border: 1px solid #10b981; padding: 10px; border-radius: 8px; font-weight: 700; text-align: center; font-size: 12px; }
    .status-badge-wait { background: rgba(245, 158, 11, 0.15); color: #fbbf24; border: 1px solid #f59e0b; padding: 10px; border-radius: 8px; font-weight: 700; text-align: center; font-size: 12px; }
    .status-badge-err { background: rgba(239, 68, 68, 0.15); color: #f87171; border: 1px solid #ef4444; padding: 10px; border-radius: 8px; font-weight: 700; text-align: center; font-size: 12px; }
    .kanban-col-header { text-align: center; font-weight: 800; padding: 12px; border-radius: 8px; margin-bottom: 15px; color: #fff; text-transform: uppercase; font-size: 14px; }
    .header-todo { background: linear-gradient(45deg, #475569, #334155); }
    .header-ip { background: linear-gradient(45deg, #f59e0b, #d97706); }
    .header-done { background: linear-gradient(45deg, #10b981, #059669); }
    </style>
""", unsafe_allow_html=True)

try:
    SUPABASE_URL = st.secrets["SUPABASE_URL"].rstrip("/")
    SUPABASE_KEY = st.secrets["SUPABASE_KEY"]
    headers = {"apikey": SUPABASE_KEY, "Authorization": f"Bearer {SUPABASE_KEY}", "Content-Type": "application/json", "Prefer": "return=representation"}
except Exception:
    st.error("Chybí konfigurace databáze.")
    st.stop()

# =========================================================================
# ZÁCHRANNÝ KRUH PŘI STISKNUTÍ F5 (Zamezí rozpadu menu)
# =========================================================================
if not st.session_state.get("prihlasen"):
    qs_user = st.query_params.get("user")
    if qs_user:
        res_auto = requests.get(f"{SUPABASE_URL}/rest/v1/uzivatele", headers=headers, params={"jmeno": f"eq.{qs_user}", "select": "*"}).json()
        if isinstance(res_auto, list) and len(res_auto) > 0:
            st.session_state.prihlasen = True
            st.session_state.role = str(res_auto[0]["role"]).lower()
            st.session_state.kredity = res_auto[0]["kredity"]
            st.session_state.uzivatel = res_auto[0]["jmeno"]
            st.session_state.skolni_kod = res_auto[0].get("skolni_kod", "")
            st.session_state.trida_nazev = res_auto[0].get("trida_nazev", "")
            st.rerun()

if not st.session_state.get("prihlasen"):
    st.warning("Spojení s aplikací bylo přerušeno (např. obnovením stránky bez parametrů).")
    if st.button("Přejít zpět na hlavní přihlašovací stránku"):
        st.switch_page("app.py")
    st.stop()

def parse_planovany_pocet(zamer_text):
    if not zamer_text: return 4
    match = re.search(r'Nahlášený počet členů týmu:\s*(\d+)', str(zamer_text))
    if match: return int(match.group(1))
    return 4

st.title("Startup Hub a Firemní Dashboard")

uzivatel = st.session_state.get("uzivatel", "")
skolni_kod = st.session_state.get("skolni_kod", "")
trida_nazev = st.session_state.get("trida_nazev", "")

# 1. Načtení OSOBNÍHO zůstatku žáka
res_akt_u = requests.get(f"{SUPABASE_URL}/rest/v1/uzivatele?jmeno=eq.{uzivatel}", headers=headers).json()
if res_akt_u and isinstance(res_akt_u, list) and len(res_akt_u) > 0:
    osobni_zustatek_zaka = int(res_akt_u[0].get("kredity", 0))
    if not skolni_kod: skolni_kod = res_akt_u[0].get("skolni_kod", "")
    if not trida_nazev: trida_nazev = res_akt_u[0].get("trida_nazev", "")
else:
    osobni_zustatek_zaka = int(st.session_state.get("kredity", 0))

# 2. Načtení firem
res_vsechny = requests.get(f"{SUPABASE_URL}/rest/v1/firmy?skolni_kod=eq.{skolni_kod}&select=*&order=id.desc", headers=headers).json()
vsechny_firmy = res_vsechny if isinstance(res_vsechny, list) else []

moje_firma = next((f for f in vsechny_firmy if uzivatel.lower() in [
    str(f.get('ceo_jmeno','')).lower(),
    str(f.get('cfo_jmeno','')).lower(),
    str(f.get('cto_jmeno','')).lower()
]), None)

if not moje_firma:
    res_z_check = requests.get(f"{SUPABASE_URL}/rest/v1/zamestnanci?jmeno_zamestnance=eq.{uzivatel}", headers=headers).json()
    if isinstance(res_z_check, list) and len(res_z_check) > 0:
        f_id_check = res_z_check[0].get('firma_id')
        moje_firma = next((f for f in vsechny_firmy if f['id'] == f_id_check), None)

if moje_firma:
    if str(moje_firma.get('ceo_jmeno', '')).lower() == uzivatel.lower(): moje_role = "CEO"
    elif str(moje_firma.get('cfo_jmeno', '')).lower() == uzivatel.lower(): moje_role = "CFO"
    elif str(moje_firma.get('cto_jmeno', '')).lower() == uzivatel.lower(): moje_role = "CTO"
    else: moje_role = "ZAMESTNANEC"
else:
    moje_role = "CEO" if st.session_state.get("role") in ["firma", "admin"] else "ZAMESTNANEC"

je_vedeni = moje_role in ["CEO", "CFO", "CTO"]

akt_dan_mtech = 15.0
akt_dan_prijem = 15.0
kurz_kc = 10.0

target_skola = moje_firma['skolni_kod'] if moje_firma else (skolni_kod or 'SYSTEM')
nastaveni_res = requests.get(f"{SUPABASE_URL}/rest/v1/skolni_nastaveni?skolni_kod=eq.{target_skola}", headers=headers).json()
if isinstance(nastaveni_res, list) and len(nastaveni_res) > 0:
    akt_dan_mtech = float(nastaveni_res[0].get('mtech_dan_pct', 15.0))
    akt_dan_prijem = float(nastaveni_res[0].get('dan_prijem_pct', 15.0))
    kurz_kc = float(nastaveni_res[0].get('kurz_kc', 10.0))

# =========================================================================
# 1. ZALOŽENÍ FIRMY (VKLAD Z OSOBNÍHO ÚČTU)
# =========================================================================
if not moje_firma:
    st.info(f"Vítejte v podnikatelském akcelerátoru M-TECH (Třída: {trida_nazev or 'Vaše třída'}).")
    st.markdown("Získali jste oprávnění k založení startupu. Vyplňte zakladatelský zápis u Notáře a odešlete spis ke schválení.")
    
    res_spoluzaci = requests.get(f"{SUPABASE_URL}/rest/v1/uzivatele?skolni_kod=eq.{skolni_kod}&trida_nazev=eq.{trida_nazev}&role=neq.ucitel", headers=headers).json()
    seznam_spoluzaku = [u['jmeno'] for u in res_spoluzaci if u['jmeno'] != uzivatel] if isinstance(res_spoluzaci, list) else []

    u_notar, u_zivnost, u_financak, u_cssz, u_rejstrik = st.tabs([
        "1. Notářský zápis", "2. Živnostenský úřad", "3. Finanční úřad", "4. ČSSZ", "5. Zápis do Rejstříku"
    ])
    
    if "reg_data" not in st.session_state: st.session_state.reg_data = {}
    
    with u_notar:
        st.session_state.reg_data["nazev_firmy"] = st.text_input("Obchodní název firmy (Startupu):", placeholder="Např. EcoTech s.r.o.")
        col_n1, col_n2, col_n3 = st.columns(3)
        with col_n1: st.session_state.reg_data["ceo"] = st.text_input("CEO (Generální ředitel):", value=uzivatel, disabled=True)
        with col_n2: st.session_state.reg_data["cfo"] = st.selectbox("CFO (Finanční ředitel):", ["-- Neobsazeno --"] + seznam_spoluzaku)
        with col_n3: st.session_state.reg_data["cto"] = st.selectbox("CTO (Technický ředitel):", ["-- Neobsazeno --"] + seznam_spoluzaku)
        st.session_state.reg_data["jednani"] = st.selectbox("Způsob jednání statutárních orgánů:", ["Každý jednatel samostatně", "Společně"])
        
        col_k1, col_k2 = st.columns(2)
        with col_k1: 
            if osobni_zustatek_zaka < 10:
                st.error(f"Váš osobní zůstatek je pouze {osobni_zustatek_zaka} M-K. Pro založení firmy potřebujete minimálně 10 M-K. Jděte do Moje peněženka -> Úřad práce!")
                vklad_hodnota = 0
            else:
                vklad_hodnota = st.number_input(
                    f"Vklad základního kapitálu z vaší peněženky (K dispozici: {osobni_zustatek_zaka} M-K):",
                    min_value=10, max_value=int(osobni_zustatek_zaka), value=min(100, int(osobni_zustatek_zaka)), step=5
                )
            st.session_state.reg_data["vklad"] = vklad_hodnota
        with col_k2: 
            st.session_state.reg_data["podily_popis"] = st.text_area("Rozdělení podílů (%):", value="CEO: 100 %")

    with u_zivnost:
        st.session_state.reg_data["druh_zivnosti"] = st.radio("Druh živnosti:", ["Volná", "Řemeslná", "Vázaná"], horizontal=True)
        col_j1, col_j2 = st.columns(2)
        with col_j1:
            st.session_state.reg_data["zivnost_detail"] = st.text_input("Obor činnosti:", placeholder="Např. 3D tisk a modelování...")
            st.session_state.reg_data["predmet"] = st.text_area("Předmět podnikání:", value="Vývoj inovativních produktů a zakázková výroba.")
        with col_j2:
            st.session_state.reg_data["bozp_garant"] = st.text_input("Odpovědný zástupce za BOZP:", value=uzivatel)
            st.session_state.reg_data["provozovna"] = st.text_input("Sídlo provozovny:", value="Školní dílna a MakerSpace")

    with u_financak:
        st.session_state.reg_data["typ_dani"] = st.multiselect("Registrace k daním:", ["DPPO (Daň z příjmu PO)", "M-TECH daň z e-shopu", "Daň ze závislé činnosti (Mzdy)"], default=["DPPO (Daň z příjmu PO)", "M-TECH daň z e-shopu"])
        st.session_state.reg_data["zdanovaci_obdobi"] = st.selectbox("Zdaňovací období:", ["Měsíční", "Čtvrtletní"])

    with u_cssz:
        st.session_state.reg_data["pocet_zakladatelu"] = st.number_input("Celkový nahlášený počet členů firmy (Vedení + Zaměstnanci):", min_value=1, max_value=20, value=4, step=1)
        st.session_state.reg_data["bozp_prohlaseni"] = st.checkbox("Čestně prohlašuji, že pracoviště splňuje bezpečnostní předpisy BOZP.", value=True)

    with u_rejstrik:
        st.session_state.reg_data["ubo"] = st.text_input("Skutečný majitel (UBO):", value=f"{uzivatel}")
        st.session_state.reg_data["kodex_souhlas"] = st.checkbox("Zavazujeme se dodržovat Etický kodex mladého podnikatele.", value=True)
        st.write("---")
        st.markdown("#### Rozšíření o reálný prodej")
        st.caption("Zapojit firmu do reálného prodeje veřejnosti (např. na Dni otevřených dveří) schvaluje výhradně učitel/Úřad — zaškrtnutím jen podáváte žádost.")
        st.session_state.reg_data["zada_urovne_3"] = st.checkbox("Žádáme o zapojení do rozšíření o reálný prodej (Úroveň 3).", value=False)
        st.write("---")
        
        vklad_k_prevodu = int(st.session_state.reg_data.get("vklad", 0))
        st.caption(f"Odesláním žádosti bude z vaší osobní peněženky převedena částka **{vklad_k_prevodu} M-K** do základního kapitálu firmy.")
        
        if st.button("Převést kapitál a odeslat spis na Kontrolní úřad", type="primary"):
            d = st.session_state.reg_data
            nazev = d.get("nazev_firmy", "").strip()
            
            if not nazev:
                st.error("Vyplňte prosím název firmy v záložce Notářský zápis.")
            elif osobni_zustatek_zaka < vklad_k_prevodu or vklad_k_prevodu < 10:
                st.error(f"Nemáte dostatek financí ve své osobní peněžence (Požadováno: {vklad_k_prevodu} M-K).")
            else:
                cfo_val = None if d.get("cfo") == "-- Neobsazeno --" else d.get("cfo")
                cto_val = None if d.get("cto") == "-- Neobsazeno --" else d.get("cto")
                
                plan_clenove = int(d.get('pocet_zakladatelu', 4))
                zpusob_jednani = d.get('jednani', 'Každý jednatel samostatně')
                
                zamer_str = f"Způsob jednání: {zpusob_jednani} | Nahlášený počet členů týmu: {plan_clenove} | Druh živnosti: {d.get('druh_zivnosti')} | Obor: {d.get('zivnost_detail', '')} | Předmět: {d.get('predmet', '')} | Garant BOZP: {d.get('bozp_garant', '')} | Provozovna: {d.get('provozovna', '')}"
                
                novy_osobni_zustatek = int(osobni_zustatek_zaka - vklad_k_prevodu)
                requests.patch(f"{SUPABASE_URL}/rest/v1/uzivatele?jmeno=eq.{uzivatel}", headers=headers, json={"kredity": novy_osobni_zustatek})
                st.session_state.kredity = novy_osobni_zustatek

                zvolena_uroven = 3 if d.get("zada_urovne_3", False) else 2
                payload = {
                    "nazev_firmy": nazev, "skolni_kod": skolni_kod, "trida_nazev": trida_nazev, "uroven_projektu": zvolena_uroven,
                    "ceo_jmeno": uzivatel, "cfo_jmeno": cfo_val, "cto_jmeno": cto_val, "podnikatelsky_zamer": zamer_str,
                    "pocatecni_kapital": int(vklad_k_prevodu), "stave_licence": "CEKA_NA_SCHVALENI", "duvod_zamitnuti": ""
                }
                res_create = requests.post(f"{SUPABASE_URL}/rest/v1/firmy", headers=headers, json=payload)
                if res_create.status_code in [200, 201]:
                    new_f_data = res_create.json()
                    new_f_id = new_f_data[0]['id'] if (isinstance(new_f_data, list) and len(new_f_data) > 0) else None
                    if new_f_id:
                        requests.post(f"{SUPABASE_URL}/rest/v1/kniha_prijmu_vydaju", headers=headers, json={
                            "firma_id": new_f_id, "typ_transakce": "PRIJEM", "titul": f"Vklad základního kapitálu zakladatelem ({uzivatel})", "castka": vklad_k_prevodu, "auditovano": True
                        })
                    st.success(f"Základní kapitál {vklad_k_prevodu} M-K byl převeden na účet firmy a žádost odeslána.")
                    st.rerun()
                else:
                    st.error(f"Chyba při zakládání firmy: {res_create.text}")
    st.stop()

# =========================================================================
# 2. VÝPOČET ZŮSTATKU FIREMNÍHO ÚČTU (Z KNIHY PŘÍJMŮ A VÝDAJŮ)
# =========================================================================
res_kniha_calc = requests.get(f"{SUPABASE_URL}/rest/v1/kniha_prijmu_vydaju?firma_id=eq.{moje_firma['id']}&select=typ_transakce,castka", headers=headers).json()
prijmy_firmy = sum(float(tx.get('castka', 0)) for tx in (res_kniha_calc if isinstance(res_kniha_calc, list) else []) if str(tx.get('typ_transakce','')).upper() in ['PRIJEM', 'PŘÍJEM'])
vydaje_firmy = sum(float(tx.get('castka', 0)) for tx in (res_kniha_calc if isinstance(res_kniha_calc, list) else []) if str(tx.get('typ_transakce','')).upper() in ['VYDEJ', 'VÝDEJ'])

if not res_kniha_calc or len(res_kniha_calc) == 0:
    firemni_zustatek = int(moje_firma.get('pocatecni_kapital', 100))
else:
    firemni_zustatek = int(prijmy_firmy - vydaje_firmy)

# Načtení zaměstnanců
res_zam_global = requests.get(f"{SUPABASE_URL}/rest/v1/zamestnanci?firma_id=eq.{moje_firma['id']}&select=*", headers=headers).json()
zamestnanci_firmy = res_zam_global if isinstance(res_zam_global, list) else []

pocet_vedeni = 1 + (1 if moje_firma.get('cfo_jmeno') else 0) + (1 if moje_firma.get('cto_jmeno') else 0)
pocet_zam = len(zamestnanci_firmy)
pocet_celkem_tym = pocet_vedeni + pocet_zam
planovany_pocet_clenu = parse_planovany_pocet(moje_firma.get('podnikatelsky_zamer', ''))

# =========================================================================
# HLAVNÍ METRIKY: DVA ODDĚLENÉ ÚČTY (FIREMNÍ vs. OSOBNÍ)
# =========================================================================
st.subheader(f"Dashboard: {moje_firma['nazev_firmy']} (Vaše pozice: {moje_role})")

col_m1, col_m2, col_m3, col_m4 = st.columns(4)
with col_m1:
    st.metric(label="Firemní pokladna (Účet firmy)", value=f"{firemni_zustatek} M-K", help="Peníze společnosti určené na výplaty, materiál, nákupy a provoz.")
with col_m2:
    st.metric(label="Můj osobní účet (Soukromé)", value=f"{osobni_zustatek_zaka} M-K", help="Vaše soukromé kapesné a výdělky.")
with col_m3:
    st.metric(label="Členové týmu", value=f"{pocet_celkem_tym} / {planovany_pocet_clenu} osob")
with col_m4:
    st.metric(label="Status rejstříku", value=moje_firma['stave_licence'])

st.write("---")

if pocet_celkem_tym < planovany_pocet_clenu:
    st.error(f"""
    **NESOULAD SE ZAKLADATELSKOU LISTINOU:** V zakladatelském spisu máte nahlášeno **{planovany_pocet_clenu} členů**, ale reálně máte zapsáno pouze **{pocet_celkem_tym} osob**.  
    V záložce **1. Zakladatelský Spis** nebo **4. Tým a HR** doplňte zbývajícího člena, jinak Kontrolní úřad neschválí licenci!
    """)

if moje_firma["stave_licence"] == "CEKA_NA_SCHVALENI":
    st.warning("Vaše firma byla odeslána na Kontrolní úřad a čeká na posouzení vyučujícím.")
    
if moje_firma["stave_licence"] == "ZAMITNUTO":
    st.error(f"""
    **ŽÁDOST BYLA KONTROLNÍM ÚŘADEM ZAMÍTNUTA / VRÁCENA K PŘEPRACOVÁNÍ**  
    **Důvod od vyučujícího:** {moje_firma.get('duvod_zamitnuti', 'Doplňte náležitosti spisu.')}  
    Upravte níže složení týmu a klikněte na tlačítko **Odeslat k novému posouzení**.
    """)

if moje_firma["stave_licence"] == "UKONCENO":
    st.error("Tato společnost byla oficiálně zlikvidována a ukončena.")

if je_vedeni:
    if moje_firma.get("stave_licence") in ["SCHVALENO", "CEKA_NA_SCHVALENI"]:
        with st.expander("📄 Zakladatelská listina (ke stažení)"):
            radky_listiny = [
                ("Název firmy", moje_firma.get("nazev_firmy", "")),
                ("Právní forma", "Studentská firma v rámci projektu M-TECH CORE"),
                ("Škola / licenční kód", moje_firma.get("skolni_kod", "")),
                ("Třída", moje_firma.get("trida_nazev", "")),
                ("Generální ředitel (CEO)", moje_firma.get("ceo_jmeno", "")),
                ("Finanční ředitel (CFO)", moje_firma.get("cfo_jmeno", "-- neobsazeno --")),
                ("Technický ředitel (CTO)", moje_firma.get("cto_jmeno", "-- neobsazeno --")),
                ("Podnikatelský záměr", moje_firma.get("podnikatelsky_zamer", "")),
                ("Zvolená úroveň projektu", f"Úroveň {moje_firma.get('uroven_projektu', 2)}" + (" (rozšíření o reálný prodej)" if int(moje_firma.get('uroven_projektu', 2)) == 3 else "")),
                ("Počáteční kapitál", f"{moje_firma.get('pocatecni_kapital', 0)} M-Kreditů"),
                ("Stav licence", moje_firma.get("stave_licence", "")),
            ]
            doc_buf = vygeneruj_formular_docx("Zakladatelská listina", radky_listiny)
            st.download_button("⬇️ Stáhnout zakladatelskou listinu (Word)", data=doc_buf, file_name=f"zakladatelska_listina_{moje_firma.get('nazev_firmy','firma')}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")    
    tab_zalozeni, tab_brand, tab_vyvoj, tab_hr, tab_kalkulace, tab_ucto, tab_burza, tab_denik, tab_likvidace = st.tabs([
        "1. Zakladatelský Spis", "2. Brand a AI Tank", "3. Řízení (Agile)", "4. Tým a HR", "5. Cenotvorba", "6. Účetnictví a Platby", "7. Burza", "8. Deník a Porady", "9. Likvidace firmy"
    ])
else:
    tab_zalozeni, tab_vyvoj, tab_denik = st.tabs([
        "1. O firmě a Můj tým", "2. Úkoly a Agilní vývoj", "3. Můj deník práce a Porady"
    ])
    tab_brand = tab_hr = tab_kalkulace = tab_ucto = tab_burza = tab_likvidace = None

res_zaci_tridy = requests.get(f"{SUPABASE_URL}/rest/v1/uzivatele?skolni_kod=eq.{skolni_kod}&trida_nazev=eq.{trida_nazev}&role=neq.ucitel", headers=headers).json()
vsechna_jmena_tridy = [u['jmeno'] for u in (res_zaci_tridy if isinstance(res_zaci_tridy, list) else [])]

# ==========================================
# ZÁLOŽKA 1: SPIS A SPRÁVA TÝMU
# ==========================================
with tab_zalozeni:
    st.subheader("Registrační spis a Právní status")
    barva_statusu = "status-ok" if moje_firma['stave_licence'] == 'SCHVALENO' else ("status-err" if moje_firma['stave_licence'] in ['UKONCENO', 'ZAMITNUTO'] else "status-wait")
    st.markdown(f"**Stav v rejstříku:** <span class='{barva_statusu}'>{moje_firma['stave_licence']}</span>", unsafe_allow_html=True)
    
    with st.container(border=True):
        st.markdown(f"**Obchodní firma:** `{moje_firma['nazev_firmy']}`")
        st.markdown(f"**Třída:** `{moje_firma.get('trida_nazev', trida_nazev)}`")
        st.markdown(f"**CEO (Generální ředitel):** `{moje_firma.get('ceo_jmeno', 'Neobsazeno')}`")
        st.markdown(f"**CFO (Finanční ředitel):** `{moje_firma.get('cfo_jmeno', 'Neobsazeno')}`")
        st.markdown(f"**CTO (Technický ředitel):** `{moje_firma.get('cto_jmeno', 'Neobsazeno')}`")
        
        if zamestnanci_firmy:
            z_seznam_txt = ", ".join([f"{z['jmeno_zamestnance']} ({z.get('pozice','Pracovník')})" for z in zamestnanci_firmy])
            st.markdown(f"**Zaměstnanci ({len(zamestnanci_firmy)}):** `{z_seznam_txt}`")
        else:
            st.markdown("**Zaměstnanci:** `Zatím žádní řadoví zaměstnanci`")
            
        st.markdown(f"**Aktuální obsazenost týmu:** `{pocet_celkem_tym} z {planovany_pocet_clenu} nahlášených osob`")
        st.markdown(f"**Základní kapitál:** `{moje_firma.get('pocatecni_kapital', 100)} M-K`")
        st.divider()
        st.markdown("**Předmět podnikání a záměr:**")
        st.write(moje_firma.get('podnikatelsky_zamer', 'Neuvedeno'))
        
        if moje_firma.get('popis'):
            st.divider()
            st.markdown("**Oficiální vyvalidovaný byznys plán z AI Mentora:**")
            st.markdown(moje_firma.get('popis'))

    if je_vedeni and moje_firma["stave_licence"] == "ZAMITNUTO":
        if st.button("Odeslat opravený spis k novému posouzení vyučujícímu", type="primary"):
            requests.patch(f"{SUPABASE_URL}/rest/v1/firmy?id=eq.{moje_firma['id']}", headers=headers, json={"stave_licence": "CEKA_NA_SCHVALENI"})
            st.success("Spis byl znovu odeslán Kontrolnímu úřadu.")
            st.rerun()

    if je_vedeni:
        st.divider()
        st.markdown("### Správa členů týmu a zaměstnanců")
        
        with st.expander("1. Statutární vedení firmy (CFO a CTO)", expanded=False):
            volni_kandidati_vedeni = ["-- Neobsazeno --"] + [j for j in vsechna_jmena_tridy if j != moje_firma.get('ceo_jmeno')]
            with st.form("form_update_vedeni_ceo_main"):
                akt_cfo = moje_firma.get('cfo_jmeno') or "-- Neobsazeno --"
                akt_cto = moje_firma.get('cto_jmeno') or "-- Neobsazeno --"
                
                idx_cfo = volni_kandidati_vedeni.index(akt_cfo) if akt_cfo in volni_kandidati_vedeni else 0
                idx_cto = volni_kandidati_vedeni.index(akt_cto) if akt_cto in volni_kandidati_vedeni else 0
                
                new_cfo = st.selectbox("CFO (Finanční ředitel):", volni_kandidati_vedeni, index=idx_cfo)
                new_cto = st.selectbox("CTO (Technický ředitel):", volni_kandidati_vedeni, index=idx_cto)
                
                if st.form_submit_button("Uložit změny vedení"):
                    patch_cfo = None if new_cfo == "-- Neobsazeno --" else new_cfo
                    patch_cto = None if new_cto == "-- Neobsazeno --" else new_cto
                    requests.patch(f"{SUPABASE_URL}/rest/v1/firmy?id=eq.{moje_firma['id']}", headers=headers, json={"cfo_jmeno": patch_cfo, "cto_jmeno": patch_cto})
                    st.success("Vedení společnosti bylo aktualizováno.")
                    st.rerun()

        with st.expander("2. Přijmout nového běžného zaměstnance do firmy", expanded=(pocet_celkem_tym < planovany_pocet_clenu)):
            obsazena_jmena = [z['jmeno_zamestnance'].lower() for z in zamestnanci_firmy]
            if moje_firma.get('ceo_jmeno'): obsazena_jmena.append(moje_firma['ceo_jmeno'].lower())
            if moje_firma.get('cfo_jmeno'): obsazena_jmena.append(moje_firma['cfo_jmeno'].lower())
            if moje_firma.get('cto_jmeno'): obsazena_jmena.append(moje_firma['cto_jmeno'].lower())
            
            volni_pro_zam = [j for j in vsechna_jmena_tridy if j.lower() not in obsazena_jmena]
            
            if not volni_pro_zam:
                st.info(f"Ve třídě {trida_nazev} jsou momentálně registrováni pouze {len(vsechna_jmena_tridy)} žáci a všichni už ve firmě mají roli.")
            else:
                with st.form("form_prijmout_zamestnance_spis_tab"):
                    col_p1, col_p2 = st.columns(2)
                    with col_p1:
                        novy_zam_jmeno = st.selectbox("Vyberte žáka k přijetí:", volni_pro_zam)
                        novy_zam_pozice = st.text_input("Pracovní pozice (např. Grafik, Operátor výroby):", value="Pracovník vývoje")
                    with col_p2:
                        novy_zam_sazba = st.number_input("Hodinová sazba (M-K / hod):", min_value=10, value=50, step=5)
                    
                    if st.form_submit_button("Přijmout zaměstnance do firmy", type="primary"):
                        if novy_zam_jmeno and novy_zam_pozice.strip():
                            requests.post(f"{SUPABASE_URL}/rest/v1/zamestnanci", headers=headers, json={
                                "firma_id": int(moje_firma["id"]), "jmeno_zamestnance": str(novy_zam_jmeno), "pozice": str(novy_zam_pozice.strip()), "hodinova_sazba": float(novy_zam_sazba), "vyplaceno_celkem": 0.0
                            })
                            st.success(f"Žák {novy_zam_jmeno} byl přijat do firmy.")
                            st.rerun()

        with st.expander("3. Úprava nahlášené kapacity týmu ve stanovách", expanded=False):
            st.caption("Pokud nemůžete sehnat dalšího spolužáka a vyučující vám povolil menší tým, upravte zde nahlášený počet.")
            with st.form("form_zmena_kapacity_stanov"):
                nova_kapacita = st.number_input("Nový celkový počet členů týmu:", min_value=1, max_value=20, value=planovany_pocet_clenu, step=1)
                if st.form_submit_button("Uložit změnu kapacity"):
                    stary_zamer = str(moje_firma.get('podnikatelsky_zamer', ''))
                    if "Nahlášený počet členů týmu:" in stary_zamer:
                        novy_zamer = re.sub(r'Nahlášený počet členů týmu:\s*\d+', f'Nahlášený počet členů týmu: {nova_kapacita}', stary_zamer)
                    else:
                        novy_zamer = f"Nahlášený počet členů týmu: {nova_kapacita} | " + stary_zamer
                    
                    requests.patch(f"{SUPABASE_URL}/rest/v1/firmy?id=eq.{moje_firma['id']}", headers=headers, json={"podnikatelsky_zamer": novy_zamer})
                    st.success("Kapacita týmu byla ve spisu upravena.")
                    st.rerun()

# ==========================================
# ZÁLOŽKA 2: BRAND A AI TANK
# ==========================================
if tab_brand:
    with tab_brand:
        t_aktiva, t_lean, t_ai_shark = st.tabs(["Vizuální Identita", "Lean Canvas", "AI Shark Tank"])
        with t_aktiva:
            with st.form("form_brand"):
                b_logo = st.text_input("Odkaz na logo:", value=moje_firma.get('logo_url','') or "")
                b_web = st.text_input("Odkaz na web:", value=moje_firma.get('web_url','') or "")
                if st.form_submit_button("Uložit odkazy"):
                    requests.patch(f"{SUPABASE_URL}/rest/v1/firmy?id=eq.{moje_firma['id']}", headers=headers, json={"logo_url": b_logo, "web_url": b_web})
                    st.rerun()

        with t_lean:
            st.info("**Tip:** Místo ručního vypisování doporučujeme použít modul **Lean Startup Validator** (v levém menu). Jakmile u AI Mentora obhájíte svůj nápad na více než 80 %, váš promyšlený Lean Canvas se sem **automaticky propíše**.")
            res_c = requests.get(f"{SUPABASE_URL}/rest/v1/lean_canvas?firma_id=eq.{moje_firma['id']}", headers=headers).json()
            exist_canvas = res_c[0] if isinstance(res_c, list) and len(res_c) > 0 else None
            with st.form("form_full_canvas"):
                c1, c2, c3 = st.columns(3)
                with c1: prob = st.text_area("1. Problém", value=exist_canvas.get("problem","") if exist_canvas else "", height=150)
                with c2: hodnota = st.text_area("2. Unikátní Hodnota", value=exist_canvas.get("hodnota","") if exist_canvas else "", height=150)
                with c3: cilovka = st.text_area("3. Cílová Skupina", value=exist_canvas.get("cilova_skupina","") if exist_canvas else "", height=150)
                c4, c5, c6 = st.columns(3)
                with c4: sol = st.text_area("4. Řešení", value=exist_canvas.get("reseni","") if exist_canvas else "", height=150)
                with c5: kanaly = st.text_area("5. Prodejní Kanály", value=exist_canvas.get("kanaly","") if exist_canvas else "", height=150)
                with c6: vyhoda = st.text_area("6. Nefér Výhoda", value=exist_canvas.get("vyhoda","") if exist_canvas else "", height=150)
                c7, c8 = st.columns(2)
                with c7: naklady = st.text_area("7. Struktura Nákladů", value=exist_canvas.get("naklady","") if exist_canvas else "", height=100)
                with c8: prijmy = st.text_area("8. Zdroje Příjmů", value=exist_canvas.get("prijmy","") if exist_canvas else "", height=100)
                if st.form_submit_button("Uložit / Přepsat Lean Canvas manuálně"):
                    c_payload = {"firma_id": moje_firma["id"], "problem": prob, "reseni": sol, "hodnota": hodnota, "cilova_skupina": cilovka, "kanaly": kanaly, "vyhoda": vyhoda, "naklady": naklady, "prijmy": prijmy}
                    if exist_canvas: requests.patch(f"{SUPABASE_URL}/rest/v1/lean_canvas?id=eq.{exist_canvas['id']}", headers=headers, json=c_payload)
                    else: requests.post(f"{SUPABASE_URL}/rest/v1/lean_canvas", headers=headers, json=c_payload)
                    st.success("Byznys plán byl manuálně uložen.")
                    st.rerun()

        with t_ai_shark:
            res_pitches = requests.get(f"{SUPABASE_URL}/rest/v1/ai_pitches?firma_id=eq.{moje_firma['id']}&order=datum.desc", headers=headers).json()
            ma_uspesnou_investici = any(p.get('schvaleno_investovano', False) for p in res_pitches) if isinstance(res_pitches, list) else False
            if ma_uspesnou_investici:
                st.success("Získali jste Seed investici od AI Shark Tanku na firemní účet.")
            else:
                with st.form("form_pitch"):
                    p_nazev = st.text_input("Název prezentace:")
                    p_popis = st.text_area("Detailní pitch:")
                    col_p1, col_p2 = st.columns(2)
                    with col_p1: p_castka = st.number_input("Požadovaný kapitál pro firmu (M-K):", min_value=50, value=200, step=10)
                    with col_p2: p_akcie = st.number_input("Nabízené akcie (ks):", min_value=5, value=20)
                    if st.form_submit_button("Spustit AI Pitching"):
                        requests.post(f"{SUPABASE_URL}/rest/v1/ai_pitches", headers=headers, json={"firma_id": moje_firma['id'], "nazev_pitchu": p_nazev, "popis_projektu": p_popis, "zadana_castka": p_castka, "nabizene_akcie": p_akcie, "hodnoceni_ostry": "[SCHVALENO]", "hodnoceni_vizionarka": "[SCHVALENO]", "hodnoceni_rychly": "[SCHVALENO]", "schvaleno_investovano": True, "investovana_castka": p_castka})
                        requests.post(f"{SUPABASE_URL}/rest/v1/kniha_prijmu_vydaju", headers=headers, json={
                            "firma_id": moje_firma["id"], "typ_transakce": "PRIJEM", "titul": f"AI Shark Tank: {p_nazev}", "castka": int(p_castka), "auditovano": True
                        })
                        st.success(f"Investice {p_castka} M-K byla připsána do firemní pokladny!")
                        st.rerun()

# ==========================================
# ZÁLOŽKA 3: AGILE
# ==========================================
if tab_vyvoj:
    with tab_vyvoj:
        with st.form("form_novy_ukol"):
            col_u1, col_u2 = st.columns([3, 1])
            with col_u1: u_nazev = st.text_input("Nový úkol pro tým:")
            with col_u2: u_sp = st.number_input("Story Points:", min_value=1, value=3)
            if st.form_submit_button("Přidat úkol do Backlogu"):
                requests.post(f"{SUPABASE_URL}/rest/v1/projektove_ukoly", headers=headers, json={"firma_id": moje_firma["id"], "nazev_ukolu": u_nazev, "zodpovedna_osoba": uzivatel, "story_points": u_sp, "stav": "TO_DO"})
                st.rerun()

        res_tasks = requests.get(f"{SUPABASE_URL}/rest/v1/projektove_ukoly?firma_id=eq.{moje_firma['id']}&order=id.asc", headers=headers).json()
        tasks = res_tasks if isinstance(res_tasks, list) else []
        col_todo, col_ip, col_done = st.columns(3)
        with col_todo:
            st.markdown("<div class='kanban-col-header header-todo'>K řešení (To Do)</div>", unsafe_allow_html=True)
            for t in [x for x in tasks if x.get('stav') == 'TO_DO']:
                with st.container(border=True):
                    st.write(f"**{t['nazev_ukolu']}** ({t.get('story_points', 1)} SP)")
                    if st.button("Začít řešit", key=f"start_{t['id']}"):
                        requests.patch(f"{SUPABASE_URL}/rest/v1/projektove_ukoly?id=eq.{t['id']}", headers=headers, json={"stav": "IN_PROGRESS"})
                        st.rerun()
        with col_ip:
            st.markdown("<div class='kanban-col-header header-ip'>V řešení (In Progress)</div>", unsafe_allow_html=True)
            for t in [x for x in tasks if x.get('stav') == 'IN_PROGRESS']:
                with st.container(border=True):
                    st.write(f"**{t['nazev_ukolu']}** ({t.get('story_points', 1)} SP)")
                    if st.button("Dokončit", key=f"done_{t['id']}"):
                        requests.patch(f"{SUPABASE_URL}/rest/v1/projektove_ukoly?id=eq.{t['id']}", headers=headers, json={"stav": "DONE"})
                        st.rerun()
        with col_done:
            st.markdown("<div class='kanban-col-header header-done'>Hotovo (Done)</div>", unsafe_allow_html=True)
            for t in [x for x in tasks if x.get('stav') == 'DONE']:
                with st.container(border=True):
                    st.write(f"**{t['nazev_ukolu']}** ({t.get('story_points', 1)} SP)")

# ==========================================
# ZÁLOŽKA 4: TÝM A HR (VÝPLATA Z FIREMNÍHO ÚČTU)
# ==========================================
if tab_hr:
    with tab_hr:
        st.markdown(f"#### Obsazenost týmu: **{pocet_celkem_tym} / {planovany_pocet_clenu} osob**")
        
        if zamestnanci_firmy:
            df_zam = pd.DataFrame(zamestnanci_firmy)
            zobrazit_sloupce = [c for c in ['jmeno_zamestnance', 'pozice', 'hodinova_sazba', 'vyplaceno_celkem'] if c in df_zam.columns]
            df_show = df_zam[zobrazit_sloupce].rename(columns={'jmeno_zamestnance': 'Jméno', 'pozice': 'Pozice', 'hodinova_sazba': 'Sazba (M-K/hod)', 'vyplaceno_celkem': 'Vyplaceno celkem'})
            st.dataframe(df_show, use_container_width=True)
        else:
            st.info("Firma zatím nemá žádné řadové zaměstnance.")

        if zamestnanci_firmy:
            st.divider()
            st.markdown("#### Výplata mezd (Hrazeno z firemní pokladny)")
            st.caption(f"Disponibilní zůstatek na firemním účtu: **{firemni_zustatek} M-K**")
            
            with st.form("form_mzdy_vyplata_tab4"):
                vybrany_z_jmeno = st.selectbox("Komu chcete vyplatit mzdu:", [z["jmeno_zamestnance"] for z in zamestnanci_firmy])
                hodiny = st.number_input("Počet odpracovaných hodin:", min_value=1.0, value=4.0, step=0.5)
                
                if st.form_submit_button("Odeslat výplatu z firemní pokladny"):
                    vybrany_z = next((z for z in zamestnanci_firmy if z["jmeno_zamestnance"] == vybrany_z_jmeno), None)
                    hruba = int(hodiny * float(vybrany_z["hodinova_sazba"]))
                    dan_castka = int(hruba * (akt_dan_prijem / 100.0))
                    cista = hruba - dan_castka
                    
                    if hruba > firemni_zustatek:
                        st.error(f"Nedostatek financí na firemním účtu! (Požadováno: {hruba} M-K, k dispozici: {firemni_zustatek} M-K)")
                    else:
                        requests.post(f"{SUPABASE_URL}/rest/v1/kniha_prijmu_vydaju", headers=headers, json={
                            "firma_id": moje_firma['id'],
                            "typ_transakce": "VYDEJ",
                            "titul": f"Výplata mzdy: {vybrany_z_jmeno} ({hodiny} hod, srážka daně {dan_castka} M-K)",
                            "castka": hruba,
                            "auditovano": True
                        })
                        
                        res_zam_ucet = requests.get(f"{SUPABASE_URL}/rest/v1/uzivatele?jmeno=eq.{vybrany_z['jmeno_zamestnance']}", headers=headers).json()
                        if res_zam_ucet:
                            novy_zam_bal = int(res_zam_ucet[0].get('kredity', 0) + cista)
                            requests.patch(f"{SUPABASE_URL}/rest/v1/uzivatele?jmeno=eq.{vybrany_z['jmeno_zamestnance']}", headers=headers, json={"kredity": novy_zam_bal})
                        
                        res_stat = requests.get(f"{SUPABASE_URL}/rest/v1/uzivatele?jmeno=eq.Stat", headers=headers).json()
                        if res_stat:
                            novy_stat_bal = int(res_stat[0].get('kredity', 0) + dan_castka)
                            requests.patch(f"{SUPABASE_URL}/rest/v1/uzivatele?jmeno=eq.Stat", headers=headers, json={"kredity": novy_stat_bal})
                        
                        requests.patch(f"{SUPABASE_URL}/rest/v1/zamestnanci?id=eq.{vybrany_z['id']}", headers=headers, json={"vyplaceno_celkem": vybrany_z.get("vyplaceno_celkem", 0) + cista})
                        
                        st.success(f"Mzda {cista} M-K byla vyplacena zaměstnanci {vybrany_z_jmeno} z firemního účtu (odvedena daň {dan_castka} M-K).")
                        st.rerun()

# ==========================================
# ZÁLOŽKA 5: CENOTVORBA
# ==========================================
if tab_kalkulace:
    with tab_kalkulace:
        with st.form("form_kalkulace"):
            prod_nazev = st.text_input("Název produktu:")
            popis = st.text_area("Popis:")
            p_naklady = st.number_input("Výrobní náklady (M-K):", min_value=0.0, value=35.0)
            marze = st.number_input("Marže (M-K):", min_value=0.0, value=50.0)
            k_cena = (p_naklady + marze) * (1 + (akt_dan_mtech / 100.0))
            st.markdown(f"**Koncová prodejní cena:** `{k_cena:.2f} M-K`")
            if st.form_submit_button("Odeslat Úřadu ke schválení"):
                requests.post(f"{SUPABASE_URL}/rest/v1/kalkulacni_listy", headers=headers, json={"firma_id": moje_firma["id"], "nazev_produktu": prod_nazev, "popis": popis, "prime_naklady": p_naklady, "marze_zisk": marze, "mtech_dan_procento": akt_dan_mtech, "konecna_cena": k_cena, "schvaleno_uradem": False})
                st.rerun()

        st.divider()
        st.markdown("#### 📋 Vytvořené kalkulační listy")
        res_moje_kalk = requests.get(f"{SUPABASE_URL}/rest/v1/kalkulacni_listy?firma_id=eq.{moje_firma['id']}&order=id.desc", headers=headers).json()
        if isinstance(res_moje_kalk, list) and len(res_moje_kalk) > 0:
            for kl in res_moje_kalk:
                stav = "✅ Schváleno" if kl.get("schvaleno_uradem") else "⏳ Čeká na schválení"
                st.write(f"**{kl['nazev_produktu']}** — {kl['konecna_cena']:.2f} M-K ({stav})")
                radky_kalk = [
                    ("Firma", moje_firma.get("nazev_firmy", "")),
                    ("Název produktu", kl.get("nazev_produktu", "")),
                    ("Popis", kl.get("popis", "")),
                    ("Přímé náklady", f"{kl.get('prime_naklady', 0)} M-K"),
                    ("Marže (odměna firmy)", f"{kl.get('marze_zisk', 0)} M-K"),
                    ("M-TECH daň", f"{kl.get('mtech_dan_procento', 0)} %"),
                    ("Konečná prodejní cena", f"{kl.get('konecna_cena', 0):.2f} M-K"),
                    ("Schváleno Úřadem", "Ano" if kl.get("schvaleno_uradem") else "Ne"),
                ]
                doc_buf_kalk = vygeneruj_formular_docx(f"Kalkulační list produktu: {kl.get('nazev_produktu','')}", radky_kalk)
                st.download_button("⬇️ Stáhnout kalkulační list (Word)", data=doc_buf_kalk, file_name=f"kalkulacni_list_{kl.get('nazev_produktu','produkt')}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", key=f"dl_kalk_{kl['id']}")
        else:
            st.info("Zatím nemáte vytvořený žádný kalkulační list.")
# ==========================================
# ZÁLOŽKA 6: ÚČETNICTVÍ A FIREMNÍ PLATBY
# ==========================================
if tab_ucto:
    with tab_ucto:
        st.subheader("Firemní bankovnictví a Platby z firemního účtu")
        st.caption(f"Aktuální disponibilní zůstatek firemní pokladny: **{firemni_zustatek} M-K**")
        
        with st.expander("Odeslat platbu z firemní pokladny (B2B, Škola, Bonus)", expanded=True):
            with st.form("form_firemni_platba_vedeni"):
                typ_f_platby = st.selectbox("Typ firemního výdaje:", [
                    "B2B platba jiné studentské firmě (Materiál / Subdodávka)",
                    "Platba škole (Pronájem prostor / Nákup materiálu ze skladu školy)",
                    "Mimořádná odměna / Bonus členovi týmu"
                ])
                
                cizi_firmy = [f for f in vsechny_firmy if f.get('id') != moje_firma['id'] and f.get('stave_licence') == 'SCHVALENO']
                clenove_naseho_tymu = [z['jmeno_zamestnance'] for z in zamestnanci_firmy]
                if moje_firma.get('ceo_jmeno'): clenove_naseho_tymu.append(moje_firma['ceo_jmeno'])
                if moje_firma.get('cfo_jmeno'): clenove_naseho_tymu.append(moje_firma['cfo_jmeno'])
                if moje_firma.get('cto_jmeno'): clenove_naseho_tymu.append(moje_firma['cto_jmeno'])
                
                if "B2B platba" in typ_f_platby:
                    cil_b2b_nazev = st.selectbox("Vyberte partnerskou firmu:", [f['nazev_firmy'] for f in cizi_firmy]) if cizi_firmy else None
                elif "Bonus" in typ_f_platby:
                    cil_bonus_jmeno = st.selectbox("Vyberte příjemce bonusu:", clenove_naseho_tymu)
                else:
                    st.info("Peníze budou poukázány na účet školy (Státní pokladna).")
                    
                castka_f_platby = st.number_input("Částka k úhradě z firemního účtu (M-K):", min_value=1, max_value=max(1, firemni_zustatek), value=min(20, max(1, firemni_zustatek)), step=1)
                popis_f_platby = st.text_input("Účel platby / Číslo faktury / Poznámka:", value="Úhrada za materiál a služby")
                
                if st.form_submit_button("Odeslat firemní platbu", type="primary"):
                    if firemni_zustatek < castka_f_platby:
                        st.error(f"Nedostatek financí ve firemní pokladně! (Máte: {firemni_zustatek} M-K)")
                    else:
                        if "B2B platba" in typ_f_platby:
                            if not cil_b2b_nazev:
                                st.error("Není vybrána žádná partnerská firma.")
                            else:
                                target_b2b = next((f for f in cizi_firmy if f['nazev_firmy'] == cil_b2b_nazev), None)
                                requests.post(f"{SUPABASE_URL}/rest/v1/kniha_prijmu_vydaju", headers=headers, json={
                                    "firma_id": moje_firma['id'], "typ_transakce": "VYDEJ", "titul": f"B2B Platba firmě {cil_b2b_nazev}: {popis_f_platby}", "castka": int(castka_f_platby), "auditovano": True
                                })
                                requests.post(f"{SUPABASE_URL}/rest/v1/kniha_prijmu_vydaju", headers=headers, json={
                                    "firma_id": target_b2b['id'], "typ_transakce": "PRIJEM", "titul": f"B2B Tržba od firmy {moje_firma['nazev_firmy']}: {popis_f_platby}", "castka": int(castka_f_platby), "auditovano": True
                                })
                                st.success(f"Platba {castka_f_platby} M-K byla převedena na účet firmy {cil_b2b_nazev}.")
                                st.rerun()
                                
                        elif "Bonus" in typ_f_platby:
                            requests.post(f"{SUPABASE_URL}/rest/v1/kniha_prijmu_vydaju", headers=headers, json={
                                "firma_id": moje_firma['id'], "typ_transakce": "VYDEJ", "titul": f"Mimořádný bonus: {cil_bonus_jmeno} ({popis_f_platby})", "castka": int(castka_f_platby), "auditovano": True
                            })
                            r_bonus_u = requests.get(f"{SUPABASE_URL}/rest/v1/uzivatele?jmeno=eq.{cil_bonus_jmeno}", headers=headers).json()
                            if r_bonus_u:
                                novy_bonus_bal = int(r_bonus_u[0].get('kredity', 0) + castka_f_platby)
                                requests.patch(f"{SUPABASE_URL}/rest/v1/uzivatele?jmeno=eq.{cil_bonus_jmeno}", headers=headers, json={"kredity": novy_bonus_bal})
                            st.success(f"Bonus {castka_f_platby} M-K byl připsán do osobní peněženky {cil_bonus_jmeno}.")
                            st.rerun()
                            
                        else:
                            requests.post(f"{SUPABASE_URL}/rest/v1/kniha_prijmu_vydaju", headers=headers, json={
                                "firma_id": moje_firma['id'], "typ_transakce": "VYDEJ", "titul": f"Platba škole / Poplatek: {popis_f_platby}", "castka": int(castka_f_platby), "auditovano": True
                            })
                            r_stat = requests.get(f"{SUPABASE_URL}/rest/v1/uzivatele?jmeno=eq.Stat", headers=headers).json()
                            if r_stat:
                                novy_stat_bal = int(r_stat[0].get('kredity', 0) + castka_f_platby)
                                requests.patch(f"{SUPABASE_URL}/rest/v1/uzivatele?jmeno=eq.Stat", headers=headers, json={"kredity": novy_stat_bal})
                            st.success(f"Platba {castka_f_platby} M-K byla uhrazena škole.")
                            st.rerun()

        st.divider()
        st.markdown("#### Kniha transakcí firmy (Příjmy a Výdaje)")
        st.caption(f"Celkové příjmy: **{prijmy_firmy:.2f} M-K** | Celkové výdaje: **{vydaje_firmy:.2f} M-K**")
        
        res_kniha = requests.get(f"{SUPABASE_URL}/rest/v1/kniha_prijmu_vydaju?firma_id=eq.{moje_firma['id']}&order=id.desc", headers=headers).json()
        if isinstance(res_kniha, list) and len(res_kniha) > 0: 
            st.dataframe(pd.DataFrame(res_kniha)[['datum', 'typ_transakce', 'titul', 'castka']], use_container_width=True)
        else:
            st.info("Kniha transakcí je prázdná.")
        if int(moje_firma.get("uroven_projektu", 2)) == 3:
            st.divider()
            st.markdown("#### 🏪 Zápis reálného prodeje (Úroveň 3)")
            st.caption("Sem zapisuj skutečné tržby z prodeje veřejnosti — v Kč. Tyto peníze jdou vždy přímo na účet školy/Unie rodičů, appka je jen eviduje pro účetní.")
            with st.form("form_realny_prodej", clear_on_submit=True):
                rp_produkt = st.text_input("Název produktu:")
                col_rp1, col_rp2 = st.columns(2)
                rp_ks = col_rp1.number_input("Počet kusů:", min_value=1, value=1)
                rp_castka = col_rp2.number_input("Celková částka (Kč):", min_value=0.0, value=0.0, step=1.0)
                rp_platba = st.selectbox("Způsob platby:", ["hotovost", "QR platba"])
                rp_poznamka = st.text_input("Poznámka (nepovinné):")
                if st.form_submit_button("Zapsat reálný prodej"):
                    if not rp_produkt.strip() or rp_castka <= 0:
                        st.error("Vyplňte název produktu a částku větší než 0.")
                    else:
                        requests.post(f"{SUPABASE_URL}/rest/v1/realny_prodej", headers=headers, json={
                            "firma_id": moje_firma["id"], "popis_produktu": rp_produkt.strip(),
                            "pocet_ks": rp_ks, "castka_kc": rp_castka, "zpusob_platby": rp_platba,
                            "zapsal": uzivatel, "poznamka": rp_poznamka.strip()
                        })
                        st.success("Reálný prodej zapsán.")
                        st.rerun()

            res_moje_rp = requests.get(f"{SUPABASE_URL}/rest/v1/realny_prodej?firma_id=eq.{moje_firma['id']}&order=datum.desc", headers=headers).json()
            if isinstance(res_moje_rp, list) and len(res_moje_rp) > 0:
                st.dataframe(pd.DataFrame(res_moje_rp)[["datum", "popis_produktu", "pocet_ks", "castka_kc", "zpusob_platby"]], use_container_width=True)
                st.metric("Celkem reálných tržeb", f"{sum(r['castka_kc'] for r in res_moje_rp):,.0f} Kč")
# ==========================================
# ZÁLOŽKA 7: BURZA
# ==========================================
if tab_burza:
    with tab_burza:
        with st.form("form_ipo"):
            pocet_akcii = st.number_input("Počet akcií k prodeji:", min_value=1, value=50)
            cena_akcie = st.number_input("Cena za 1 akcii (M-K):", min_value=1.0, value=15.0)
            if st.form_submit_button("Zveřejnit na burze"):
                requests.post(f"{SUPABASE_URL}/rest/v1/burza_nabidky", headers=headers, json={"firma_id": moje_firma["id"], "pocet_k_prodeji": pocet_akcii, "cena_za_kus": cena_akcie, "aktivni": True})
                st.success("Akcie zveřejněny.")
                st.rerun()

# ==========================================
# ZÁLOŽKA 8: DENÍK A PORADY
# ==========================================
if tab_denik:
    with tab_denik:
        t_osobni, t_porady = st.tabs(["Individuální výkazy", "Zápisy z firemních porad"])
        
        with t_osobni:
            st.markdown("Zde evidujte, na čem jste konkrétně pracovali.")
            with st.form("form_denik_prace_osobni"):
                dp_popis = st.text_area("Co přesně jste udělali / odpracovali:")
                dp_hodiny = st.number_input("Počet odpracovaných hodin:", min_value=0.5, max_value=12.0, value=1.0, step=0.5)
                if st.form_submit_button("Uložit do osobního výkazu"):
                    if dp_popis.strip():
                        requests.post(f"{SUPABASE_URL}/rest/v1/denik_prace", headers=headers, json={
                            "jmeno_zaka": uzivatel, "firma_id": moje_firma['id'], "popis_prace": dp_popis.strip(), "hodiny": dp_hodiny
                        })
                        st.success("Záznam byl úspěšně uložen.")
                        st.rerun()
                    else:
                        st.error("Vyplňte popis práce.")
            
            st.markdown("#### Historie odvedené práce ve firmě")
            res_denik = requests.get(f"{SUPABASE_URL}/rest/v1/denik_prace?firma_id=eq.{moje_firma['id']}&order=id.desc", headers=headers).json()
            if isinstance(res_denik, list):
                bezne_vykazy = [p for p in res_denik if not str(p.get("popis_prace", "")).startswith("[ZÁPIS Z PORADY]")]
                if bezne_vykazy:
                    df_denik = pd.DataFrame(bezne_vykazy)
                    zobrazit = [c for c in ['datum', 'jmeno_zaka', 'popis_prace', 'hodiny'] if c in df_denik.columns]
                    df_show = df_denik[zobrazit].rename(columns={'datum': 'Datum', 'jmeno_zaka': 'Pracovník', 'popis_prace': 'Popis činnosti', 'hodiny': 'Hodiny'})
                    st.dataframe(df_show, use_container_width=True)
                else:
                    st.info("Zatím nebyly zaznamenány žádné pracovní výkazy.")

        with t_porady:
            st.markdown("Zde evidujte oficiální zápisy z vašich firemních schůzek a porad.")
            if je_vedeni:
                with st.form("form_zapis_porady"):
                    st.caption(f"Zápis podává: **{uzivatel}** (Role: **{moje_role}**)")
                    zapis_text = st.text_area("Co se řešilo, jaké jsou úkoly a závěry z porady:", height=150)
                    if st.form_submit_button("Uložit zápis z porady"):
                        if zapis_text.strip():
                            requests.post(f"{SUPABASE_URL}/rest/v1/denik_prace", headers=headers, json={
                                "jmeno_zaka": f"{uzivatel} ({moje_role})", "firma_id": moje_firma['id'], "popis_prace": f"[ZÁPIS Z PORADY]\n{zapis_text.strip()}", "hodiny": 0
                            })
                            st.success("Zápis z porady byl uložen.")
                            st.rerun()
                        else:
                            st.error("Zápis nesmí být prázdný.")
            else:
                st.info("Zápisy z firemních porad smí do systému vkládat pouze členové vedení. Zaměstnanci mohou zápisy číst.")
            
            st.markdown("#### Historie porad")
            if isinstance(res_denik, list):
                porady_list = [p for p in res_denik if str(p.get("popis_prace", "")).startswith("[ZÁPIS Z PORADY]")]
                if porady_list:
                    for p in porady_list:
                        with st.container(border=True):
                            st.caption(f"**Datum:** {p.get('datum', '')[:10]} | **Zapsal:** {p.get('jmeno_zaka', '')}")
                            cisty_text = p.get('popis_prace', '').replace("[ZÁPIS Z PORADY]\n", "")
                            st.write(cisty_text)
                else:
                    st.info("Firma zatím nemá žádné zápisy z porad.")

# ==========================================
# ZÁLOŽKA 9: LIKVIDACE A UKONČENÍ FIRMY
# ==========================================
if tab_likvidace:
    with tab_likvidace:
        st.subheader("Ukončení a likvidace společnosti")
        with st.container(border=True):
            st.markdown("#### Právní kroky likvidace:")
            st.markdown("1. **Vypořádání závazků:** Vyhodnocení virtuálních M-Kreditů a zaplacení virtuální daně.")
            st.markdown("2. **Rozdělení likvidačního zůstatku:** Zbývající M-Kredity zůstávají zakladatelům firmy.")
            st.markdown("3. **Výmaz z rejstříku:** Oznámení Kontrolnímu úřadu o definitivním ukončení činnosti.")
        st.divider()
        st.markdown("#### Rozdělení dividendy (nepovinné, před ukončením)")
        st.caption("Rozděl část virtuálního zisku firmy mezi vlastníky akcií (investory), pokud firma nějaké akcie prodala.")

        res_akcionari = requests.get(f"{SUPABASE_URL}/rest/v1/vlastnici_akcii?firma_id=eq.{moje_firma['id']}&select=*", headers=headers).json()
        akcionari = res_akcionari if isinstance(res_akcionari, list) else []
        celkem_prodanych_akcii = sum(int(a.get('pocet_akcii', 0)) for a in akcionari)

        if celkem_prodanych_akcii == 0:
            st.info("Firma zatím neprodala žádné akcie, dividendu není komu vyplatit.")
        else:
            st.write(f"Firma má {celkem_prodanych_akcii} prodaných podílů mezi {len(akcionari)} investory.")
            dividenda_pct = st.slider("Kolik % aktuálního zůstatku rozdělit jako dividendu:", 0, 50, 20, key="divid_pct")
            castka_na_rozdeleni = firemni_zustatek * (dividenda_pct / 100.0)
            st.write(f"K rozdělení: **{castka_na_rozdeleni:.1f} M-K** (podle podílu na {celkem_prodanych_akcii} akciích)")

            if st.button("Vyplatit dividendu akcionářům", key="btn_vyplat_dividendu"):
                for a in akcionari:
                    podil = int(a.get('pocet_akcii', 0)) / celkem_prodanych_akcii
                    castka_akcionari = castka_na_rozdeleni * podil
                    if castka_akcionari > 0:
                        res_akcionar_u = requests.get(f"{SUPABASE_URL}/rest/v1/uzivatele?jmeno=eq.{a['majitel_jmeno']}", headers=headers).json()
                        if res_akcionar_u:
                            requests.patch(f"{SUPABASE_URL}/rest/v1/uzivatele?jmeno=eq.{a['majitel_jmeno']}", headers=headers, json={
                                "kredity": res_akcionar_u[0]['kredity'] + castka_akcionari
                            })
                requests.post(f"{SUPABASE_URL}/rest/v1/kniha_prijmu_vydaju", headers=headers, json={
                    "firma_id": moje_firma['id'], "typ_transakce": "VYDEJ",
                    "titul": f"Výplata dividendy akcionářům ({dividenda_pct}% zisku)", "castka": castka_na_rozdeleni, "auditovano": True
                })
                st.success("Dividenda byla vyplacena.")
                st.rerun()

        st.divider()       
        st.markdown("#### 📄 Protokol o likvidaci a finančním vypořádání")
        res_realny_lik = requests.get(f"{SUPABASE_URL}/rest/v1/realny_prodej?firma_id=eq.{moje_firma['id']}", headers=headers).json()
        celkem_realnych_kc = sum(float(r.get('castka_kc', 0)) for r in (res_realny_lik if isinstance(res_realny_lik, list) else []))
        radky_protokol = [
            ("Firma", moje_firma.get("nazev_firmy", "")),
            ("Škola / třída", f"{moje_firma.get('skolni_kod','')} / {moje_firma.get('trida_nazev','')}"),
            ("Úroveň projektu", f"Úroveň {moje_firma.get('uroven_projektu', 2)}"),
            ("Zůstatek firemního účtu (virtuálně)", f"{firemni_zustatek} M-Kreditů"),
            ("Reálný výtěžek z prodeje veřejnosti", f"{celkem_realnych_kc:,.0f} Kč (zůstává ve Fondu rozvoje Unie rodičů/školy)"),
            ("Stav likvidace", moje_firma.get("stave_licence", "")),
            ("Datum vystavení", pd.Timestamp.now().strftime("%d.%m.%Y")),
        ]
        doc_buf_lik = vygeneruj_formular_docx("Protokol o likvidaci a finančním vypořádání", radky_protokol)
        st.download_button("⬇️ Stáhnout likvidační protokol (Word)", data=doc_buf_lik, file_name=f"likvidacni_protokol_{moje_firma.get('nazev_firmy','firma')}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        st.divider()
        if moje_firma.get("stave_licence") == "UKONCENO":
            st.info("Společnost již byla úspěšně zlikvidována a vymazána z rejstříku.")
        else:
            with st.form("form_likvidace_firmy"):
                duvod = st.text_area("Důvod ukončení činnosti / zpráva likvidátora pro Kontrolní úřad:")
                souhlas = st.checkbox("Prohlašuji, že veškeré závazky firmy byly vyřešeny.", value=False)
                
                if st.form_submit_button("Definitivně ukončit firmu a podat žádost o výmaz", type="primary"):
                    if souhlas and duvod.strip():
                        requests.patch(f"{SUPABASE_URL}/rest/v1/firmy?id=eq.{moje_firma['id']}", headers=headers, json={
                            "stave_licence": "UKONCENO",
                            "duvod_zamitnuti": f"[LIKVIDACE] {duvod.strip()}"
                        })
                        st.success("Společnost byla úspěšně ukončena.")
                        st.rerun()
                    else:
                        st.error("Pro ukončení musíte vyplnit důvod a potvrdit vyřešení závazků.")
